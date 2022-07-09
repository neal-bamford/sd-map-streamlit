from docx import Document
from docx.shared import Inches

doc = Document('addImage.docx')
tables = doc.tables

p = tables[0].rows[0].cells[0].add_paragraph()
r = p.add_run()

r.add_picture('resized.png',width=Inches(4.0), height=Inches(.7))

p = tables[1].rows[0].cells[0].add_paragraph()

r = p.add_run()

r.add_picture('teste.png',width=Inches(4.0), height=Inches(.7))

doc.save('addImage.docx')