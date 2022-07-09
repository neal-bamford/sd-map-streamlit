from __future__ import print_function
from mailmerge import MailMerge
from datetime import date
from docx import Document
from docx.shared import Inches
import base64
import shutil

template = "../templates/Practical-Business-Python-image.docx"

## Take original templage and add all images
# Pre-process to add images
doc = Document(template)

# Reference the tables. This is where we will add our images
tables = doc.tables
p = tables[0].rows[0].cells[0].add_paragraph()
r = p.add_run()
# r.add_picture("../docx_generation/images/plot-example.jpg",width=Inches(4.0), height=Inches(.7))
r.add_picture("../docx_generation/images/plot-example.jpg")

## Save the template and reference it for the merge to happen in the next part
template = "../docx_generation/Practical-Business-Python-image-template.docx"
doc.save(template)


# Merge text
document = MailMerge(template)
# print(document.get_merge_fields())

s = document.merge(
    status='Gold',
    city='Springfield',
    phone_number='800-555-5555',
    Business='Cool Shoes',
    zip='55555',
    purchases='$500,00000000',
    shipping_limit='$5',
    state='MO',
    address='1234 Main Street',
    date='{:%d-%b-%Y}'.format(date.today()),
    discount='5%',
    recipient='Mr. Jones')

## Save the template
document.write("../docx_generation/test-output.docx")

## Read it into an object, this will be passed back in the response in Streamlit
for file in ["../docx_generation/test-output.docx"]:
        with open(file, "rb") as docx:
            encoded = base64.b64encode(docx.read()).decode()

