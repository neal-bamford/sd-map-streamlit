from __future__ import print_function

import base64
from datetime import date
import os
import shutil
import sys
import time
import uuid
import mailmerge as mm

from docx import Document
import docx
from docx.shared import Inches
import folium
from fpdf import FPDF
from selenium import webdriver


from selenium.webdriver.chrome.service import Service as ChromeService

from lib import masters_data_analytics_lib as mlib
from lib import file_tools as ft

import numpy as np
import pandas as pd
import streamlit as st


reports_generation_clean_temp_files = st.secrets["reports_generation"]["clean_temp_file"]

def stats(df, borough, ward_name, oacode, column):
    """
    Generate stats from a standardised data frame
    """
    oacode_lvl_sum  = df.loc[df["OAcode"] == oacode][column].values[0]
    
    ward_lvl_sum    = df.loc[(df["WARD_NAME"] == ward_name) \
                           & (df["borough"] == borough)].agg({column: ["sum"]}).values[0][0]
    
    borough_lvl_sum = df.loc[(df["borough"] == borough)].agg({column: ["sum"]}).values[0][0]
    
    city_lvl_mean   = df.groupby("borough").agg({column: ["sum"]}).mean().round(0).values[0].astype(int)
    
    city_lvl_sum    = df.groupby("borough").agg({column: ["sum"]}).sum().round(0).values[0].astype(int)

    return {
     "oacode_sum" : oacode_lvl_sum,
       "ward_sum" : ward_lvl_sum,
    "borough_sum" : borough_lvl_sum,
   "borough_mean" : city_lvl_mean,
       "city_sum" : city_lvl_sum
   }

def ranking(df, borough, column, limit=5):
    """
    Generate ranking at borough level for a return first, last and requested stats
    """
    ## Create a reduced dataframe
    ranking_df = df[["WARD_NAME", "borough", column]]
    ranking_df_stats_ward = ranking_df.groupby(["borough", "WARD_NAME"]).agg(total=pd.NamedAgg(column=column, aggfunc="sum"))

    ## Don't use the column as an index
    ranking_df_stats_borough = ranking_df.groupby(["borough"], as_index=False).agg(total=pd.NamedAgg(column=column, aggfunc="sum"))
    ranking_df_stats_borough["rank"] = ranking_df_stats_borough["total"].rank(method="max", ascending=False)
    ranking_df_stats_borough = ranking_df_stats_borough.sort_values(by=["rank"], ascending=True)

    return ranking_df_stats_borough.head(limit), ranking_df_stats_borough.loc[ranking_df_stats_borough["borough"] == borough], ranking_df_stats_borough.tail(limit)

def series_format(post_codes):
    """
    Comma separate a list of post codes
    """
    ret_str = ""
    for post_code in post_codes:
        ret_str += post_code + ", "
        
    ret_str = ret_str[0:len(ret_str)-2] 
    
    replacement = " and "
    reverse_replacement = replacement[::-1]
    
    ret_str = ret_str[::-1].replace(" ,", reverse_replacement, 1)[::-1]

    return ret_str

## City
city = "london"


## Data Files Locations
sd_london_postcodes_file = "./data/streamlit_{}_postcodes_oa.csv".format(city)
sd_london_population_oa_file = "./data/streamlit_{}_population_oa.csv".format(city)
sd_london_household_oa_file = "./data/streamlit_{}_household_population_oa.csv".format(city)
sd_london_qualification_oa_file = "./data/streamlit_{}_qualifictation_population_oa.csv".format(city)

###
### This method loads the data required for this page's functionality. Maybe it should be moved somehwere
### else so that the data might be shared across multiple pages
###
@st.cache
def load_data(sd_london_postcodes_file, sd_london_population_oa_file, sd_london_household_oa_file, sd_london_qualification_oa_file):
    ##
    ## LOAD ALL THE DATA FILES
    ##
    ## London Post Codes
    sd_london_postcodes_df = mlib.csv_to_dataframe(sd_london_postcodes_file)
    
    ## London Population
    sd_london_population_oa_df   = mlib.csv_to_dataframe(sd_london_population_oa_file)
    
    ## Household Data
    sd_london_household_oa_df   = mlib.csv_to_dataframe(sd_london_household_oa_file)
    
    ## Education Data
    sd_london_qualification_oa_df   = mlib.csv_to_dataframe(sd_london_qualification_oa_file)
    
    return sd_london_postcodes_df, sd_london_population_oa_df, sd_london_household_oa_df, sd_london_qualification_oa_df

##
## This method is used to create a hyperlink to download the generated report by
##
def create_download_link(val, filename):
    b64 = base64.b64encode(val)  # val looks like b'...'
    return f'<a href="data:application/octet-stream;base64,{b64.decode()}" download="{filename}">Download file</a>'
    
    
## This should load the data then the return cached so it won't do it again until ilvalidated on purpouse.
## I don't know if there's a timeout function for the cache    
sd_london_postcodes_df, sd_london_population_oa_df, sd_london_household_oa_df, sd_london_qualification_oa_df=load_data(sd_london_postcodes_file, sd_london_population_oa_file, sd_london_household_oa_file, sd_london_qualification_oa_file)

##
## Start of Streamlit
##
st.markdown("# SD General Report")
st.sidebar.markdown("# SD General Report")

## Input the post code to search with
post_code_search = st.text_input("Post Code")
 
## Capture the input and run generate_report
generate_report_link = st.button("Generate General Report")


##
## Generate Link
##
if generate_report_link:
    ##
    ## START THE STATS
    ##
    ## Generate a session id - this will come from Streamlit and will segment users
    session_id = str(uuid.uuid4())[:8]
    
    try:
        OAcode = sd_london_postcodes_df.loc[sd_london_postcodes_df["Post_Code"] == post_code_search]["OAcode"].values[0]
    except:
        raise Exception("Unable to find post code {}".format(post_code_search))
    
    other_post_codes = sd_london_postcodes_df.loc[sd_london_postcodes_df["OAcode"] == OAcode]["Post_Code"].to_numpy()
    other_post_codes = np.delete(other_post_codes, np.where(other_post_codes == post_code_search))
    number_of_boroughs = len(sd_london_postcodes_df["borough"].unique())
    
    ## Obtain the ward_name from the OACode
    ward_name = sd_london_postcodes_df.loc[sd_london_postcodes_df["OAcode"] == OAcode]["WARD_NAME"].values[0]
    borough   = sd_london_postcodes_df.loc[sd_london_postcodes_df['OAcode'] == OAcode]["borough"].values[0]
    
    other_wards = sd_london_postcodes_df.loc[sd_london_postcodes_df["borough"] == borough]["WARD_NAME"]
    other_wards = other_wards.unique()
    # print(type(other_wards)) 
    # other_wards = other_wards.to_numpy()
    other_wards = np.delete(other_wards, np.where(other_wards == borough))
    
    ###
    ### FOLIUM MAP
    ###
    if 1==1:
        
        
        ###
        ### Set path to Geckodriver.exe for local development
        ###
        
        ## check for local gecko driver
        # try: 
        #     app_path = os.path.join(st.secrets.gecko_driver.gecko_driver_path)
        #     os.environ["PATH"] += os.pathsep + app_path
        # except Exception:
        #     pass

        post_code_search_longitude = sd_london_postcodes_df.loc[sd_london_postcodes_df["Post_Code"] == post_code_search]["longitude"].to_numpy()
        post_code_search_latitude = sd_london_postcodes_df.loc[sd_london_postcodes_df["Post_Code"] == post_code_search]["latitude"].to_numpy()
        
        pc_longitudes = sd_london_postcodes_df.loc[sd_london_postcodes_df["borough"] == borough]["longitude"].to_numpy()
        pc_latitudes  = sd_london_postcodes_df.loc[sd_london_postcodes_df["borough"] == borough]["latitude"].to_numpy()
        
        pc_longitudes_max = pc_longitudes.max()
        pc_longitudes_min = pc_longitudes.min()
        
        pc_latitudes_max = pc_latitudes.max()
        pc_latitudes_min = pc_latitudes.min()
        
        pc_longitudes_mid = (pc_longitudes_max + pc_longitudes_min)/2
        pc_latitudes_mid  = (pc_latitudes_max + pc_latitudes_min)/2
        
        print(pc_longitudes_mid)
        print(pc_latitudes_mid)
        
        print("post_code_search:{}".format(post_code_search))
        print("other_post_codes:{}".format(other_post_codes))
        print("OAcode:{}".format(OAcode))
        print("ward_name:{}".format(ward_name))
        print("borough:{}".format(borough))
        
        combined = np.column_stack((pc_latitudes, pc_longitudes))
        # print(type(pc_longitudes))
        # print((pc_longitudes).head(5))
        # print(type(pc_latitudes))
        
         
        ## Create the map and add borough overlay = centre on centre of map 
        m = folium.Map(location=[pc_latitudes_mid, pc_longitudes_mid], zoom_start=12)
        
        m.fit_bounds([[pc_latitudes_min, pc_longitudes_min], [pc_latitudes_max, pc_longitudes_max]])
        
        m.get_root().html.add_child(folium.Element("""
        <style>
        .mapText {
            white-space: nowrap;
            color:blue;
            font-weight: bold;
            font-size:large
        }
        </style>"""))
        
        ## Add Ward Name Icon
        # folium.Marker([pc_latitudes_mid, pc_longitudes_mid], popup=popup,
        #                      icon=folium.Icon(icon_color='green')).add_to(m)
                             
        ## Add Ward Name Name                     
        folium.Marker([pc_latitudes_mid, pc_longitudes_mid],
                            icon=folium.DivIcon(html="" + ward_name + "",
                            class_name="mapText")).add_to(m)
                                                               
        folium.Rectangle(combined, color="green", 
                            weight=2,
                            fill=True,
                            fill_color="pink",
                            fill_opacity=0.2).add_to(m)
        
        ## Search Post_Code
        tooltip = post_code_search
        folium.Marker([post_code_search_latitude, post_code_search_longitude], 
                            icon=folium.DivIcon(html="" + post_code_search + "",
                            class_name="mapText")).add_to(m)
        
        ## Name the html file to be created with a {path}/{session_id}_map_{city}_{borough}_{ward_name}_{post_code_search}
        file = "./reports/generation/images/{}_map_{}_{}_{}_{}".format(session_id, city, borough, ward_name, post_code_search)
        file = file.replace(" ", "_").lower()
        
        ## Save the HTML file
        html_file = file + ".html"
        m.save(html_file)
        
        
        

        # options = webdriver.ChromeOptions()
        # options.set_capability("loggingPrefs", {'performance': 'ALL'})
        # options.add_argument("--headless")
        # options.add_argument("-no-sandbox")
        # options.add_argument("--disable-dev-shm-usage")
        #
        # service = ChromeService(executable_path = st.secrets.firefox.binary_location)
        # # open it, go to a website, and get results
        # browser = webdriver.Chrome(service=service
        #                           ,options=options
        #                           ,executable_path=st.secrets.gecko.binary_location)
        # wd.get("https://www.website.com")
        
        
        ## Load it
        # options = Options()
        # options.add_argument("--headless")
        # options.headless = True
        ## Set the location of FireFox
        # options.binary_location = st.secrets.firefox.binary_location
        # browser = webdriver.Firefox(options=options
        #                             ## Set the location of Geckodriver
        #                         , executable_path=st.secrets.gecko.binary_location)
        
        
        # firefox_binary = FirefoxBinary('/usr/bin/firefox/')
        # driver = webdriver.Firefox(firefox_binary=firefox_binary)
        #
        # firefoxOptions = Options()
        # firefoxOptions.add_argument("--headless")
        # service = Service(GeckoDriverManager().install())
        # browser = webdriver.Firefox(
        #     options=firefoxOptions,
        #     service=service
        # )
        
        
        from selenium.webdriver.chrome.options import Options as ChromiumOptions
        from selenium.webdriver.chrome.service import Service as ChromeService
        
        chrome_options = ChromiumOptions()
        chrome_options.headless = True
        
        chrome_options.add_argument("--headless")
        chrome_options.add_argument("-no-sandbox")
        chrome_options.add_argument("--disable-dev-shm-usage")

        chrome_service = ChromeService(executable_path = st.secrets.chrome.binary_location)
        browser = webdriver.Chrome(service=chrome_service,
                                   options = chrome_options) 
                
        ## Read in the file we created previously
        browser.get("file:///" + os.path.abspath(html_file))
        
        ## Wait for the browser to complete
        time.sleep(st.secrets.selenium.browser_pause_s)
        
        ## Save the map graphic
        location_png_file = file + ".png"
        browser.save_screenshot(location_png_file)
        browser.quit()
        
    ###
    ### Location
    ###
    if 1==1: 
        print("\n")
        print("###")
        print("### Location")
        print("###")
        print("\n")
        print("\n")
        
        print("post code    :{}".format(post_code_search))
        print("OAcode       :{}".format(OAcode))
        print("ward name    :{}".format(ward_name))
        print("borough      :{}".format(borough))
        print("----------------")
        print("\n")
        
        location_field_01 = ("The post code {} belongs to the ward {} and borough {} within the city of {}. There {} {} other post code{} which the following data is part of. {} being {}. There are {} other ward{} in the borough which {} {}.").format(
                                                                                         post_code_search,
                                                                                         ward_name, 
                                                                                         borough, 
                                                                                         city.capitalize(), 
                                                                                         ("are" if len(other_post_codes) > 1 else "is"), 
                                                                                         len(other_post_codes), 
                                                                                         ("s" if len(other_post_codes) > 1 else ""), 
                                                                                         ("These" if len(other_post_codes) > 1 else "This"), 
                                                                                         series_format(other_post_codes), 
                                                                                         len(other_wards), 
                                                                                         ("s" if len(other_wards) > 1 else ""), 
                                                                                         ("are" if len(other_wards) > 1 else "is"), 
                                                                                         series_format(other_wards)
                                                                                         )
        
        print(location_field_01)
        # print("\n")
        # print("\n")
    
    ###
    ### Population
    ###
    if 1==1: 
        print("\n")
        print("###")
        print("### Population")
        print("###")
        print("\n")
        print("\n")
        
        pop_all_stats     = stats(sd_london_population_oa_df, borough, ward_name, OAcode, "All")
        pop_male_stats    = stats(sd_london_population_oa_df, borough, ward_name, OAcode, "Males")
        pop_female_stats  = stats(sd_london_population_oa_df, borough, ward_name, OAcode, "Females")
        pop_density_stats = stats(sd_london_population_oa_df, borough, ward_name, OAcode,  "DensityPPH")

        ### Obtain rankings for this compared to others at the borough level
        limit = 5
        top, this, bottom = ranking(sd_london_population_oa_df, borough, "DensityPPH", limit)
    
        print("Population Density Ranking")
        print("==========================")
    
        ### Stats for the searched for borough    
        population_field_01 = "The population density of {} is ranked {:g} of {} at {:.2f}." \
             .format(borough, this["rank"].values[0], number_of_boroughs, round(this["total"].values[0], 2))
        print(population_field_01)
        
        
        population_field_02 = "Which is {} the average borough population density of {:.2f}.".format("above" if round(this["total"].values[0], 2) > pop_density_stats["borough_mean"] else "below", pop_density_stats["borough_mean"])
        print(population_field_02)
        
        ### If it's not the first then display the first
        population_field_03 = ""
        if this["rank"].values[0] != 1:
            population_field_03 = "{} has the highest population density at {:.2f}." \
                 .format(top.iloc[0]["borough"], round(top.iloc[0]["total"]), 2)
            print(population_field_03)
        
        population_field_04 = ""
        if this["rank"].values[0] != number_of_boroughs:
            population_field_04 ="{} has the lowest population density at {:.2f}." \
                 .format(bottom.iloc[-1]["borough"], round(bottom.iloc[-1]["total"]), 2)
            print(population_field_04)
            
        ### Male female ratio
        pop_male_female_borough_total = pop_male_stats["borough_sum"] + pop_female_stats["borough_sum"]
        pop_male_ratio = round(pop_male_stats["borough_sum"]/pop_male_female_borough_total * 100,0)
        pop_female_ratio = round(pop_female_stats["borough_sum"]/pop_male_female_borough_total * 100,0)
        
        pop_male_female_city_borough_total = pop_male_stats["city_sum"] + pop_female_stats["city_sum"]
        pop_male_city_ratio = round(pop_male_stats["city_sum"]/pop_male_female_city_borough_total * 100,0)
        pop_female_city_ratio = round(pop_female_stats["city_sum"]/pop_male_female_city_borough_total * 100,0)
        
        ### Testing Data for below
        # pop_male_ratio = 51
        # pop_female_ratio = 49
        # pop_male_city_ratio = 51
        # pop_female_city_ratio = 49
        # print(pop_male_ratio)
        # print(pop_female_ratio)
       
        ### Simple text formatting funcion
        def hls_str(r1, r2):
            return "higher than" if r1 > r2 else "lower than" if r1 < r2 else "the same as"
        
        population_field_05 = ""
        ### What to print
        if pop_male_ratio > pop_female_ratio:
            population_field_05 = "Males account for {:g}% of the borough population, which is {} the average of {:g}% at borough level. Females account for {:g}% which is {} the average of {:g}% at borough level."\
                 .format(pop_male_ratio, hls_str(pop_male_ratio, pop_male_city_ratio), \
                         pop_male_city_ratio, pop_female_ratio, hls_str(pop_female_ratio, pop_female_city_ratio), pop_female_city_ratio)
                 
        elif pop_male_ratio < pop_female_ratio:
            population_field_05 = "Females account for {:g}% of the borough population, which is {} the average of {:g}% at borough level. Males account for {:g}% which is {} the average of {:g}% at borough level."\
                 .format(pop_female_ratio, hls_str(pop_female_ratio, pop_female_city_ratio),\
                         pop_female_city_ratio, pop_male_ratio, hls_str(pop_male_ratio, pop_male_city_ratio), pop_male_city_ratio)
        else:
            population_field_05 ="Males and females are equal for the borough. The borough level average is males {:g}% and females {:g}%."\
                 .format(pop_male_city_ratio, pop_female_city_ratio)
            
        print(population_field_05)
        
        ### Combine for the report
        population_field_01 = population_field_01 + " " + population_field_02 + " " + population_field_03 + " " + population_field_04 + " " + population_field_05        
        # print("\n")
        # print("\n")
    
    ###
    ### Household
    ###
    if 1==1:
        print("\n")
        print("\n")
        print("###")
        print("### Household")
        print("###")
        print("\n")
        print("\n")
        
        hous_commerical_stats = stats(sd_london_household_oa_df, borough, ward_name, OAcode, "CommercialBuilding")
        hous_detatched_stats  = stats(sd_london_household_oa_df, borough, ward_name, OAcode, "Detached")
        hous_flat_stats       = stats(sd_london_household_oa_df, borough, ward_name, OAcode, "Flat")
        hous_semi_stats       = stats(sd_london_household_oa_df, borough, ward_name, OAcode, "Semi_detached")
        hous_terraced_stats   = stats(sd_london_household_oa_df, borough, ward_name, OAcode, "Terraced")
        
        ## Create a data frame of the findings
        household_location_stats_columns = ["household_type", "d_nd", "oacode_sum", "ward_sum", "borough_sum", "city_borough_mean", "city_sum" ]
        household_location_stats_data    = [
                                            [   "commercial", "nd", hous_commerical_stats["oacode_sum"], hous_commerical_stats["ward_sum"], hous_commerical_stats["borough_sum"], hous_commerical_stats["borough_mean"], hous_commerical_stats["city_sum"]],
                                            [     "detached",  "d",  hous_detatched_stats["oacode_sum"],  hous_detatched_stats["ward_sum"],  hous_detatched_stats["borough_sum"],  hous_detatched_stats["borough_mean"],  hous_detatched_stats["city_sum"]],
                                            [         "flat",  "d",       hous_flat_stats["oacode_sum"],       hous_flat_stats["ward_sum"],       hous_flat_stats["borough_sum"],       hous_flat_stats["borough_mean"],       hous_flat_stats["city_sum"]],
                                            ["semi-detached",  "d",       hous_semi_stats["oacode_sum"],       hous_semi_stats["ward_sum"],       hous_semi_stats["borough_sum"],       hous_semi_stats["borough_mean"],       hous_semi_stats["city_sum"]],
                                            [     "terraced",  "d",   hous_terraced_stats["oacode_sum"],   hous_terraced_stats["ward_sum"],   hous_terraced_stats["borough_sum"],   hous_terraced_stats["borough_mean"],   hous_terraced_stats["city_sum"]]
                                           ]
        
        
        
        
        household_location_stats_df = pd.DataFrame(household_location_stats_data, columns=household_location_stats_columns)
        
        ## Remove the non dwelling data i.e. Commercial buildings
        ## And sort descending
        household_location_d_stats_df = household_location_stats_df.loc[household_location_stats_df["d_nd"] == "d"]
        household_location_d_stats_df = household_location_d_stats_df.sort_values(by=["oacode_sum"], ascending=False)
    
        household_type_pretty = {"detached":"Detached", "flat":"Flat", "semi-detached":"Semi Detached", "terraced":"Terraced"}
        
        print("Residential dwellings at the OAcode level are ranked")
        print("====================================================")
        house_hold = []
        for i in range(0, 4):
            str = "{} - {} - OAcode:{} - ward:{} - borough:{} - borough avg:{}".format(i + 1, \
                                                                                       household_type_pretty[household_location_d_stats_df.iloc[i]["household_type"]], \
                                                                                       household_location_d_stats_df.iloc[i]["oacode_sum"], \
                                                                                       household_location_d_stats_df.iloc[i]["ward_sum"],   \
                                                                                       household_location_d_stats_df.iloc[i]["borough_sum"],   \
                                                                                       household_location_d_stats_df.iloc[i]["city_borough_mean"],   \
                                                                                      )
            
            str2 = "{} - OAcode:{} - ward:{} - borough:{} - borough avg:{}".format(household_type_pretty[household_location_d_stats_df.iloc[i]["household_type"]], \
                                                                                   household_location_d_stats_df.iloc[i]["oacode_sum"], \
                                                                                   household_location_d_stats_df.iloc[i]["ward_sum"],   \
                                                                                   household_location_d_stats_df.iloc[i]["borough_sum"],   \
                                                                                   household_location_d_stats_df.iloc[i]["city_borough_mean"],   \
                                                                                  )
            print(str)
            house_hold.append(str2)
            
            
    ###
    ### Education
    ###
    if 1==1:
        print("\n")
        print("\n")
        print("###")
        print("### Education")
        print("###")
        print("\n")
        print("\n")
        
        edu_unknown_stats = stats(sd_london_qualification_oa_df, borough, ward_name, OAcode, "UnkownQualification")
        edu_none_stats    = stats(sd_london_qualification_oa_df, borough, ward_name, OAcode, "NoQualification")
        edu_level1_stats  = stats(sd_london_qualification_oa_df, borough, ward_name, OAcode, "Level1")
        edu_level2_stats  = stats(sd_london_qualification_oa_df, borough, ward_name, OAcode, "Level2")
        edu_level3_stats  = stats(sd_london_qualification_oa_df, borough, ward_name, OAcode, "Level3")
        edu_level4_stats  = stats(sd_london_qualification_oa_df, borough, ward_name, OAcode, "Level4")
        edu_other_stats   = stats(sd_london_qualification_oa_df, borough, ward_name, OAcode, "OtherQualifications")
        
        ## Create a data frame of the findings
        qualification_location_stats_columns = ["qualification_type", "oacode_sum", "ward_sum", "borough_sum", "city_borough_mean", "city_sum" ]
        qualification_location_stats_data    = [
                                                [    "no_qualification",           edu_none_stats["oacode_sum"],   edu_none_stats["ward_sum"],   edu_none_stats["borough_sum"],   edu_none_stats["borough_mean"],   edu_none_stats["city_sum"]],
                                                [              "level1",         edu_level1_stats["oacode_sum"], edu_level1_stats["ward_sum"], edu_level1_stats["borough_sum"], edu_level1_stats["borough_mean"], edu_level1_stats["city_sum"]],
                                                [              "level2",         edu_level2_stats["oacode_sum"], edu_level2_stats["ward_sum"], edu_level2_stats["borough_sum"], edu_level2_stats["borough_mean"], edu_level2_stats["city_sum"]],
                                                [              "level3",         edu_level3_stats["oacode_sum"], edu_level3_stats["ward_sum"], edu_level3_stats["borough_sum"], edu_level3_stats["borough_mean"], edu_level3_stats["city_sum"]],
                                                [              "level4",         edu_level4_stats["oacode_sum"], edu_level4_stats["ward_sum"], edu_level4_stats["borough_sum"], edu_level4_stats["borough_mean"], edu_level4_stats["city_sum"]],
                                                ["other_qualifications",          edu_other_stats["oacode_sum"],  edu_other_stats["ward_sum"],  edu_other_stats["borough_sum"],  edu_other_stats["borough_mean"],  edu_other_stats["city_sum"]]
                                               ]

        qualification_location_stats_df = pd.DataFrame(qualification_location_stats_data, columns=qualification_location_stats_columns)
        
        ## Sort descending
        qualification_location_stats_df = qualification_location_stats_df.sort_values(by=["oacode_sum"], ascending=False)
    
        education_type_pretty = {"no_qualification":"No qualifications", "level1":"Level 1", "level2":"Level 2", "level3":"Level 3", "level4":"Level 4", "other_qualifications":"Other qualifications"}
        
        print("Qualifications at the OAcode level are ranked")
        print("=============================================")
        education = []
        for i in range(0, 6):
            str = "{} - {} - OAcode:{} - ward:{} - borough:{} - borough avg:{}".format(i + 1, \
                                                                                       education_type_pretty[qualification_location_stats_df.iloc[i]["qualification_type"]], \
                                                                                       qualification_location_stats_df.iloc[i]["oacode_sum"], \
                                                                                       qualification_location_stats_df.iloc[i]["ward_sum"],   \
                                                                                       qualification_location_stats_df.iloc[i]["borough_sum"],   \
                                                                                       qualification_location_stats_df.iloc[i]["city_borough_mean"]
                                                                                      )
            
            str2 = "{} - OAcode:{} - ward:{} - borough:{} - borough avg:{}".format(education_type_pretty[qualification_location_stats_df.iloc[i]["qualification_type"]], \
                                                                                   qualification_location_stats_df.iloc[i]["oacode_sum"], \
                                                                                   qualification_location_stats_df.iloc[i]["ward_sum"],   \
                                                                                   qualification_location_stats_df.iloc[i]["borough_sum"],   \
                                                                                   qualification_location_stats_df.iloc[i]["city_borough_mean"]
                                                                                   )
            print(str)
            education.append(str2)
           
    ## Start creating the report for the template
    template_name = "sd_general_report_processor_template.docx"
    stage_01_template = "./reports/templates/{}".format(template_name);
    
    ## Take original templage and add all images
    # Pre-process to add images
    image_document = Document(stage_01_template)
    
    # Reference the tables. This is where we will add our images
    image_tables = image_document.tables
    paragraph = image_tables[0].rows[0].cells[0].add_paragraph()
    paragraph_run = paragraph.add_run()
    paragraph.alignment = 1

    # r.add_picture("../docx_generation/images/plot-example.jpg",width=Inches(4.0), height=Inches(.7))
    paragraph_run.add_picture(location_png_file,width=Inches(6.0))

    ## These may move to a list then be called at the end as a cleanup process    
    ft.remove_temp_file(html_file, reports_generation_clean_temp_files)
    ft.remove_temp_file(location_png_file, reports_generation_clean_temp_files)

    
    #
    # ## Save the template and reference it for the merge to happen in the next part
    stage_02_template = "./reports/generation/documents/{}_stage_02_template_{}".format(session_id, template_name)
    image_document.save(stage_02_template)
    
    
    # Merge text
    merge_document = mm.MailMerge(stage_02_template)
    
    merge_document.merge(
        post_code_search    = post_code_search,
        borough             = borough,
        city                = city.capitalize(),
        location_field_01   = location_field_01,
        population_field_01 = population_field_01,
        house_hold_01       = house_hold[0],
        house_hold_02       = house_hold[1],
        house_hold_03       = house_hold[2],
        house_hold_04       = house_hold[3],
        education_01        = education[0],
        education_02        = education[1],
        education_03        = education[2],
        education_04        = education[3],
        education_05        = education[4],
        education_06        = education[5],
    )
    
    ## Save the template
    stage_03_docx = "./reports/generation/documents/{}_{}".format(session_id, template_name)
    merge_document.write(stage_03_docx)

    ## Remove the intermediate template    
    merge_document.close()
    ft.remove_temp_file(stage_02_template, reports_generation_clean_temp_files)
    
    ### 
    ### Read in the document to send out on the link
    ###    
    for file in [stage_03_docx]:
            with open(file, "rb") as docx:
                encoded = docx.read()
    
    html = create_download_link(encoded , template_name)
    
    st.markdown(html, unsafe_allow_html=True)