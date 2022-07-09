from __future__ import print_function
from fpdf import FPDF
from mailmerge import MailMerge
from datetime import date
from docx import Document
from docx.shared import Inches
import streamlit as st
import base64
import shutil

st.markdown("# Word Report Page")
st.sidebar.markdown("# Word Report Page")

report_text = st.text_input("Report Text")
 

export_as_word = st.button("Export Report")

def create_download_link(val, filename):
    b64 = base64.b64encode(val)  # val looks like b'...'
    return f'<a href="data:application/octet-stream;base64,{b64.decode()}" download="{filename}.docx">Download file</a>'

if export_as_word:

    template = "./templates/Practical-Business-Python-image.docx"
    
    ## Take original templage and add all images
    # Pre-process to add images
    doc = Document(template)
    
    # Reference the tables. This is where we will add our images
    tables = doc.tables
    p = tables[0].rows[0].cells[0].add_paragraph()
    r = p.add_run()
    # r.add_picture("../docx_generation/images/plot-example.jpg",width=Inches(4.0), height=Inches(.7))
    r.add_picture("./docx_generation/images/plot-example.jpg")
    
    ## Save the template and reference it for the merge to happen in the next part
    template = "./docx_generation/Practical-Business-Python-image-template.docx"
    doc.save(template)
    
    
    # Merge text
    document = MailMerge(template)
    # print(document.get_merge_fields())
    
    s = document.merge(
        status='Gold',
        city='Springfield',
        phone_number='800-555-5555',
        Business='Cool Shoes',
        zip='55555',
        purchases='$500,00000000',
        shipping_limit='$5',
        state='MO',
        address='1234 Main Street',
        date='{:%d-%b-%Y}'.format(date.today()),
        discount='5%',
        recipient='Mr. Jones')
    
    ## Save the template
    output_docx = "./docx_generation/test-output.docx"
    document.write(output_docx)
    
    for file in [output_docx]:
            with open(file, "rb") as docx:
                encoded = docx.read()
    
    html = create_download_link(encoded , "test-output")

    st.markdown(html, unsafe_allow_html=True)