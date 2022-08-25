from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches
from docx.shared import Pt       #Helps to specify font size
from docx.shared import RGBColor
from docx.shared import RGBColor #Helps to specify font Color
from fpdf import FPDF
from lib import file_tools as ft
from pathlib import Path
from unidecode import unidecode

import docx
import json
import logging
import mailmerge as mm
import os
import pandas as pd
import sys

log = logging.getLogger(__name__)


def map_alignment(alignment):
  log.debug(f"alignment:{alignment}")
  if alignment.lower() == "center" or alignment.lower() == "centre":
    return WD_TABLE_ALIGNMENT.CENTER
  elif alignment.lower() == "right":
    return WD_TABLE_ALIGNMENT.RIGHT
  else:
    return WD_TABLE_ALIGNMENT.LEFT

def map_text_alignment(text_alignment):
  # log.debug(f"text_alignment:{text_alignment}")
  
  if (text_alignment.lower() == "center") or (text_alignment.lower() == "centre"):
    # log.debug(f"Matched Center")
    return WD_PARAGRAPH_ALIGNMENT.CENTER
  elif (text_alignment.lower() == "right") or (text_alignment.lower() == "RIGHT (2)".lower()):
    # log.debug(f"Matched Right")
    return WD_PARAGRAPH_ALIGNMENT.RIGHT
  elif (text_alignment.lower() == "justify") or (text_alignment.lower() == "JUSTIFY (3)".lower()):
    # log.debug(f"Matched Jusity")
    return WD_PARAGRAPH_ALIGNMENT.JUSTIFY
  else:
    # log.debug(f"Default Left")
    return WD_PARAGRAPH_ALIGNMENT.LEFT
  
def clear_cell():
  """
  Clear the contents of the cell retrieved from globals[]["table_cell"]
  """
  if "table_cell" not in globals():
    log.error("Can not reference table_cell")
  else:
    table_cell = globals()["table_cell"]
    table_cell.text = ""
    
def delete_row():
  """
  Deletes the row of the table if we choose not to add the item
  """
  if "table_cell" not in globals():
    log.error("Can not reference table_cell")
  else:
    table_cell = globals()["table_cell"]
    parent = table_cell._parent
    
    row_to_delete = None
    for row in parent.rows:
      for cell in row.cells:
        if cell._element == table_cell._element:
          row_to_delete = row
    
    if row_to_delete != None:
      parent._tbl.remove(row_to_delete._tr)
  
def include_image(image_name_in_context, image_alignment="left", image_width=6):
  if "table_cell" not in globals():
    log.error("Can not reference table_cell")
  else:
    table_cell = globals()["table_cell"]
    image_location = globals()["report_context"][image_name_in_context]
    table_cell.text = ""
    paragraph = table_cell.add_paragraph()
    paragraph_run = paragraph.add_run()
    paragraph.alignment = map_alignment(image_alignment)
    paragraph_run.add_picture(image_location, width=Inches(image_width))
  
def include_text(text_content, text_alignment='left', format_tokens=None):
  """
  Add text to a table cell - copies the attributes of the text font name, color and size 
  and re-applies after changing the text. 
  """

  if "table_cell" not in globals():
    log.error("Can not reference table_cell")
  else:
    table_cell = globals()["table_cell"]
    text_content = globals()["report_context"][text_content]
    
    ##format if format_tokens
    if format_tokens != None:
      text_content = text_content.replace("{{", "{")
      text_content = text_content.replace("}}", "}")
      text_content = text_content.format(format_tokens)

    ## Copy the for name, size and color before changing the text
    font_name = table_cell.paragraphs[0].runs[0].font.name
    font_size = table_cell.paragraphs[0].runs[0].font.size
    font_color = table_cell.paragraphs[0].runs[0].font.color
    ## take existing alignment and map it
    
    ## Change the text
    table_cell.text = text_content
    
    ## Set the alignment if we pass it in
    if text_alignment != None:
      log.debug(f"text_alignment:{text_alignment}")
      alignment = map_text_alignment(str(text_alignment))
      table_cell.paragraphs[0].alignment = alignment

    ## Reference the new
    font = table_cell.paragraphs[0].runs[0].font
    ## Set the values
    font.name = font_name
    font.size = font_size
    font.color.rgb = font_color.rgb

def include_error(error_content, text_alignment='left'):
  """
  Adds any error encountered to the table cell
  """
  if "table_cell" not in globals():
    log.error("Can not reference table_cell")
  else:
    table_cell = globals()["table_cell"]
    paragraph = table_cell.add_paragraph()
    paragraph.alignment = map_alignment(text_alignment)
    run = paragraph.add_run(error_content)
    run.font.color.rgb = RGBColor(255, 0, 0)

def include_table(data, shading=None, style=None, columns=None):
  """
  Include a Pandas DataFrame as a Word Table
  Apply a style to it - the style MUST exist in the template document
  Shading as a [[],[]] list of lists ignore the columns
  reorder the columns with columns
  """
  if "table_cell" not in globals():
    log.error("Can not reference table_cell")
  else:
    table_cell = globals()["table_cell"]
    
    ## Reference available styles
    document_styles = table_cell._parent._parent._parent.styles

    ## Make a note of what's there, we need to remove it later
    existing_paragraphs = table_cell.paragraphs

    ## If we want shading then make a dataframe with the 
    ## original data column
    if shading != None:
        shading_df = pd.DataFrame(shading, columns=data.columns.tolist())
        
    ## make a frame from the columns - should be able to re-order here too
    if columns != None:
        data = data[columns]
        
    ## Create the new table to populate with our data
    dataframe_table = table_cell.add_table(data.shape[0]+1, data.shape[1])

    ## Remove the existing paragraph so it fits properly
    for paragraph in existing_paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    ## Needed?
    dataframe_table.allow_autofit = True
    dataframe_table.autofit = True

    ## Add any table style - the table style MUST exist in the word doc being used
    if style != None:
      for style in document_styles:
        log.debug(style.name)

      dataframe_table.style = style
        

    ## Populate the table with our dataframe
    
    ### add the header rows.
    for j in range(data.shape[-1]):
        dataframe_table.cell(0,j).text = data.columns[j]

    ### add the rest of the data frame
    for i in range(data.shape[0]):
        for j in range(data.shape[-1]):
            dataframe_table.cell(i+1,j).text = str(data.values[i,j])    

    ## Add any cell colouring
    if shading != None:
        
        ## If we want to change the columns then
        if columns != None:
            data = data[columns]
            shading_df = shading_df[columns]
        
        ## We might have re-ordered the columns so remake the list
        shading = shading_df.values.tolist()
        
        s_r_idx = 0                          # Style row index
        for shading_row in shading:          # Style rows LOOP
            s_c_idx = 0                      # Style cell index
            for shading_cell in shading_row: # Style cells LOOP
                if shading_cell:             # There's a shading in there, let's use it!
                    ## Remove # in colour
                    shading_cell = shading_cell.replace("#", "")
                    ## Create the shading element and apply it
                    shading_element = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), shading_cell))
                    dataframe_table.rows[s_r_idx+1].cells[s_c_idx]._tc.get_or_add_tcPr().append(shading_element)
                s_c_idx += 1
            
            s_r_idx += 1
            
def generate_report(session_id
                  , report_context
                  , properties) -> str:
  """
  Generate the word report based on a merge and a replacement with Python scripting...
  """
  stage_03_docx = "remove_me_once_working_properly"
  globals()["report_context"] = report_context
  locals().update(report_context)
  
  ### LOAD THE TEMPLATE PROCESSOR
  template_processor_file_name = report_context["template_processor_file_name"]
  with open(template_processor_file_name) as template_processor_file:
      template_processor = json.load(template_processor_file)
      
  json_obj_to_use = None
  report_option = report_context["report_option"]
  
  ### LOAD THE REPORT OPTION, DEFAULT TO FIRST IF ONLY 1
  for json_obj_idx in template_processor["report_options"]:
    
    ## Always set the json object to use to the first
    if json_obj_to_use == None or json_obj_idx["report_option"] == report_option:
      json_obj_to_use = json_obj_idx
    
  template_absolute_path = json_obj_to_use["report_template"]
  template_name = os.path.basename(template_absolute_path)
  report_generation_folder = json_obj_to_use["report_generation_folder"]
      
  remove_temp_files =  properties["reports_generation"]["clean_temp_files"]        

  ## Start creating the report for the template
  # template_absolute_path = "sd_general_report_processor_template.docx"
  stage_01_template = template_absolute_path
  
  ## Take original template and add all images
  # Pre-process to add images
  table_document = Document(stage_01_template)

  ###
  ### ADD IMAGES AND TEXT TO TABLE CELLS
  ###     
  # Reference the tables. This is where we will add our images
  tables = table_document.tables
  
  image_files_to_remove = []
  ## Loop through the json
  indexes = {}
  # locals = {}

  ## Loop through the tables, rows and cells and do what is necessary
  for table in tables:
    for row in table.rows:
      for cell in row.cells:
        table_cell_text = cell.text
        
        if table_cell_text.startswith("{{"):
          # Add the table cell to the globals
          globals()["table_cell"] = cell ##Might need a UUID to segment session
        
          # Take the value found that we think is a command and strip out the command part
          table_cell_text = table_cell_text.replace("{{", "").replace("}}", "")
          
          table_cell_text = unidecode(table_cell_text)
          
          # Call the code and catch the response
          try:
            response = exec(table_cell_text, globals() ,locals())
          except Exception as response_error:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            log.error(exc_type, exc_obj, exc_tb)
            include_error(str(response_error))
        else:
          log.warn("Not a Command")            
    
    ## Save the template and reference it for the merge to happen in the next part
    stage_02_template = "{}/{}_stage_02_template_{}".format(report_generation_folder, session_id, template_name)
    table_document.save(stage_02_template)
    
    
    # Merge text
    merge_document = mm.MailMerge(stage_02_template)
    merge_document.merge(**report_context)
    
    ## Save the template
    stage_03_docx = "{}/{}_{}".format(report_generation_folder, session_id, template_name)
    merge_document.write(stage_03_docx)
    
    ## Remove the intermediate template    
    merge_document.close()
    
    ## These may move to a list then be called at the end as a cleanup process    
    ## remove files generated here
    ft.remove_temp_file(stage_02_template, remove_temp_files)
    
    ## remove image files we referenced
    if remove_temp_files:
        for image_file_name in image_files_to_remove:
            ft.remove_temp_file(image_file_name, remove_temp_files)
  
  return stage_03_docx