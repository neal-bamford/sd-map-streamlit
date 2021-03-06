from docx import Document
from pathlib import Path
import mailmerge as mm
import docx
import os
import json
from docx.shared import Inches
from fpdf import FPDF
from lib import file_tools as ft

def generate_report(session_id
                  , report_context
                  , properties):
    
    ### LOAD THE TEMPLATE PROCESSOR
    template_processor_file_name = report_context["template_processor_file_name"]
    with open(template_processor_file_name) as template_processor_file:
        template_processor = json.load(template_processor_file)


    template_absolute_path = template_processor["report_template"]
    template_name = os.path.basename(template_absolute_path)
    report_generation_folder = template_processor["report_generation_folder"]
        
    remove_temp_files =  properties["reports_generation"]["clean_temp_files"]        

    ## Start creating the report for the template
    # template_absolute_path = "sd_general_report_processor_template.docx"
    stage_01_template = template_absolute_path
    
    ## Take original templage and add all images
    # Pre-process to add images
    image_document = Document(stage_01_template)

    ###
    ### ADD IMAGES
    ###     
    # Reference the tables. This is where we will add our images
    image_tables = image_document.tables
    
    image_files_to_remove = []
    ## Loop through the json
    indexes = {}
    for image in template_processor["images"]:
        include = image["include"]
        
        ## Check it's in the report context
        if image["image_ref"] in report_context:
          image_file_name = report_context[image["image_ref"]]
  
          ## If the image file exists add it to the files to remove
          ## regardless of the include flag        
          if Path(image_file_name).is_file:
              image_files_to_remove.append(image_file_name)
          
          if include:
              paragraph = image_tables[image["table_cell"]].rows[0].cells[0].add_paragraph()
              paragraph_run = paragraph.add_run()
              paragraph.alignment = image["display_align"]
              paragraph_run.add_picture(image_file_name, width=Inches(image["image_width_inch"]))

    #
    # ## Save the template and reference it for the merge to happen in the next part
    stage_02_template = "{}/{}_stage_02_template_{}".format(report_generation_folder, session_id, template_name)
    image_document.save(stage_02_template)
    
    
    # Merge text
    merge_document = mm.MailMerge(stage_02_template)
    merge_document.merge(**report_context)
    
    ## Save the template
    stage_03_docx = "{}/{}_{}".format(report_generation_folder, session_id, template_name)
    merge_document.write(stage_03_docx)
    
    ## Remove the intermediate template    
    merge_document.close()
    
    ## These may move to a list then be called at the end as a cleanup process    
    ## remove files generated here
    ft.remove_temp_file(stage_02_template, remove_temp_files)
    
    ## remove image files we referenced
    if remove_temp_files:
        for image_file_name in image_files_to_remove:
            ft.remove_temp_file(image_file_name, remove_temp_files)
    
    return stage_03_docx